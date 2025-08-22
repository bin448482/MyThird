#!/usr/bin/env python3
"""
简历文档解析器
支持多种文档格式的解析：Markdown、Word、PDF、纯文本
"""

import os
import logging
from pathlib import Path
from typing import Dict, Optional
try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False
    magic = None
import markdown
from docx import Document
import PyPDF2
import pdfplumber

from .exceptions import (
    ResumeProcessingError, UnsupportedFormatError, DocumentParseError,
    FileOperationError, handle_exception, log_exception
)

logger = logging.getLogger(__name__)


class ResumeDocumentParser:
    """简历文档解析器"""
    
    SUPPORTED_FORMATS = ['md', 'docx', 'pdf', 'txt']
    
    def __init__(self, config: Dict = None):
        """
        初始化解析器
        
        Args:
            config: 配置字典，包含解析选项
        """
        self.config = config or {}
        self.max_file_size_mb = self.config.get('max_file_size_mb', 10)
        self.encoding = self.config.get('encoding', 'utf-8')
        
        logger.info(f"初始化简历文档解析器，支持格式: {self.SUPPORTED_FORMATS}")
    
    def detect_format(self, file_path: str) -> str:
        """
        检测文档格式
        
        Args:
            file_path: 文件路径
            
        Returns:
            str: 文档格式 ('md', 'docx', 'pdf', 'txt')
            
        Raises:
            UnsupportedFormatError: 不支持的文档格式
        """
        if not os.path.exists(file_path):
            raise DocumentParseError(f"文件不存在: {file_path}")
        
        # 检查文件大小
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        if file_size_mb > self.max_file_size_mb:
            raise DocumentParseError(f"文件过大: {file_size_mb:.1f}MB > {self.max_file_size_mb}MB")
        
        # 首先通过文件扩展名判断
        file_ext = Path(file_path).suffix.lower()
        ext_mapping = {
            '.md': 'md',
            '.markdown': 'md',
            '.docx': 'docx',
            '.doc': 'docx',  # 将.doc也当作docx处理
            '.pdf': 'pdf',
            '.txt': 'txt'
        }
        
        if file_ext in ext_mapping:
            detected_format = ext_mapping[file_ext]
            logger.debug(f"通过扩展名检测到格式: {detected_format}")
            return detected_format
        
        # 如果扩展名不明确，使用magic库检测MIME类型（如果可用）
        if MAGIC_AVAILABLE:
            try:
                mime_type = magic.from_file(file_path, mime=True)
                logger.debug(f"检测到MIME类型: {mime_type}")
                
                mime_mapping = {
                    'text/markdown': 'md',
                    'text/x-markdown': 'md',
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
                    'application/msword': 'docx',
                    'application/pdf': 'pdf',
                    'text/plain': 'txt'
                }
                
                if mime_type in mime_mapping:
                    detected_format = mime_mapping[mime_type]
                    logger.debug(f"通过MIME类型检测到格式: {detected_format}")
                    return detected_format
            except Exception as e:
                logger.warning(f"MIME类型检测失败: {e}")
        else:
            logger.debug("magic库不可用，跳过MIME类型检测")
        
        # 默认当作文本文件处理
        logger.warning(f"无法确定文件格式，默认当作文本文件处理: {file_path}")
        return 'txt'
    
    def extract_content(self, file_path: str) -> str:
        """
        提取文档内容（统一入口）
        
        Args:
            file_path: 文件路径
            
        Returns:
            str: 提取的文本内容
            
        Raises:
            DocumentParseError: 文档解析失败
        """
        try:
            format_type = self.detect_format(file_path)
            logger.info(f"开始解析文档: {file_path} (格式: {format_type})")
            
            if format_type == 'md':
                content = self.parse_markdown(file_path)
            elif format_type == 'docx':
                content = self.parse_word_document(file_path)
            elif format_type == 'pdf':
                content = self.parse_pdf(file_path)
            elif format_type == 'txt':
                content = self.parse_text_file(file_path)
            else:
                raise UnsupportedFormatError(f"不支持的文档格式: {format_type}")
            
            if not content or not content.strip():
                raise DocumentParseError(f"文档内容为空: {file_path}")
            
            logger.info(f"文档解析完成，内容长度: {len(content)} 字符")
            return content.strip()
            
        except Exception as e:
            logger.error(f"文档解析失败: {file_path}, 错误: {e}")
            raise DocumentParseError(f"文档解析失败: {e}")
    
    def parse_markdown(self, file_path: str) -> str:
        """
        解析Markdown文档
        
        Args:
            file_path: Markdown文件路径
            
        Returns:
            str: 解析后的纯文本内容
        """
        try:
            with open(file_path, 'r', encoding=self.encoding) as f:
                md_content = f.read()
            
            # 将Markdown转换为HTML，然后提取纯文本
            html = markdown.markdown(md_content)
            
            # 简单的HTML标签清理
            import re
            # 移除HTML标签
            text = re.sub(r'<[^>]+>', '', html)
            # 处理HTML实体
            text = text.replace('&nbsp;', ' ')
            text = text.replace('&lt;', '<')
            text = text.replace('&gt;', '>')
            text = text.replace('&amp;', '&')
            
            # 清理多余的空白字符
            text = re.sub(r'\n\s*\n', '\n\n', text)
            text = re.sub(r' +', ' ', text)
            
            return text.strip()
            
        except UnicodeDecodeError:
            # 尝试其他编码
            for encoding in ['gbk', 'gb2312', 'latin1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        md_content = f.read()
                    html = markdown.markdown(md_content)
                    text = re.sub(r'<[^>]+>', '', html)
                    return text.strip()
                except UnicodeDecodeError:
                    continue
            raise DocumentParseError(f"无法解码Markdown文件: {file_path}")
        except Exception as e:
            raise DocumentParseError(f"Markdown解析失败: {e}")
    
    def parse_word_document(self, file_path: str) -> str:
        """
        解析Word文档
        
        Args:
            file_path: Word文档路径
            
        Returns:
            str: 解析后的纯文本内容
        """
        try:
            doc = Document(file_path)
            
            # 提取所有段落文本
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(' | '.join(row_text))
            
            content = '\n\n'.join(paragraphs)
            return content.strip()
            
        except Exception as e:
            raise DocumentParseError(f"Word文档解析失败: {e}")
    
    def parse_pdf(self, file_path: str) -> str:
        """
        解析PDF文档
        
        Args:
            file_path: PDF文件路径
            
        Returns:
            str: 解析后的纯文本内容
        """
        content = ""
        
        # 首先尝试使用pdfplumber（更好的文本提取）
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        content += page_text + "\n\n"
            
            if content.strip():
                logger.debug("使用pdfplumber成功提取PDF内容")
                return content.strip()
        except Exception as e:
            logger.warning(f"pdfplumber解析失败，尝试PyPDF2: {e}")
        
        # 如果pdfplumber失败，尝试PyPDF2
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        content += page_text + "\n\n"
            
            if content.strip():
                logger.debug("使用PyPDF2成功提取PDF内容")
                return content.strip()
            else:
                raise DocumentParseError("PDF内容提取为空")
                
        except Exception as e:
            raise DocumentParseError(f"PDF解析失败: {e}")
    
    def parse_text_file(self, file_path: str) -> str:
        """
        解析纯文本文件
        
        Args:
            file_path: 文本文件路径
            
        Returns:
            str: 文件内容
        """
        try:
            with open(file_path, 'r', encoding=self.encoding) as f:
                content = f.read()
            return content.strip()
            
        except UnicodeDecodeError:
            # 尝试其他编码
            for encoding in ['gbk', 'gb2312', 'latin1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    logger.debug(f"使用编码 {encoding} 成功读取文件")
                    return content.strip()
                except UnicodeDecodeError:
                    continue
            raise DocumentParseError(f"无法解码文本文件: {file_path}")
        except Exception as e:
            raise DocumentParseError(f"文本文件解析失败: {e}")
    
    def validate_content(self, content: str) -> bool:
        """
        验证提取的内容是否有效
        
        Args:
            content: 提取的文本内容
            
        Returns:
            bool: 内容是否有效
        """
        if not content or not content.strip():
            return False
        
        # 检查内容长度
        if len(content.strip()) < 50:
            logger.warning("提取的内容过短，可能不是有效的简历")
            return False
        
        # 检查是否包含常见的简历关键词
        resume_keywords = [
            '姓名', '电话', '邮箱', '教育', '工作', '经验', '技能', '项目',
            'name', 'phone', 'email', 'education', 'work', 'experience', 'skills', 'project'
        ]
        
        content_lower = content.lower()
        keyword_count = sum(1 for keyword in resume_keywords if keyword in content_lower)
        
        if keyword_count < 2:
            logger.warning("内容中缺少常见的简历关键词")
            return False
        
        return True
    
    def get_document_info(self, file_path: str) -> Dict:
        """
        获取文档基本信息
        
        Args:
            file_path: 文件路径
            
        Returns:
            Dict: 文档信息
        """
        try:
            file_stat = os.stat(file_path)
            format_type = self.detect_format(file_path)
            
            return {
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'format': format_type,
                'size_bytes': file_stat.st_size,
                'size_mb': file_stat.st_size / (1024 * 1024),
                'modified_time': file_stat.st_mtime,
                'is_supported': format_type in self.SUPPORTED_FORMATS
            }
        except Exception as e:
            logger.error(f"获取文档信息失败: {e}")
            return {
                'file_path': file_path,
                'error': str(e)
            }


def create_resume_parser(config: Dict = None) -> ResumeDocumentParser:
    """
    创建简历文档解析器的便捷函数
    
    Args:
        config: 配置字典
        
    Returns:
        ResumeDocumentParser: 解析器实例
    """
    return ResumeDocumentParser(config)


# 示例用法
if __name__ == "__main__":
    # 配置日志
    logging.basicConfig(level=logging.INFO)
    
    # 创建解析器
    parser = ResumeDocumentParser({
        'max_file_size_mb': 10,
        'encoding': 'utf-8'
    })
    
    # 示例：解析文档
    try:
        # content = parser.extract_content("path/to/resume.pdf")
        # print(f"提取的内容长度: {len(content)}")
        # print(f"内容预览: {content[:200]}...")
        print("简历文档解析器初始化完成")
    except Exception as e:
        print(f"解析失败: {e}")